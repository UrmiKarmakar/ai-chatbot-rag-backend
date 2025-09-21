# documents/models.py
import os
import logging
from django.db import models
from django.contrib.auth import get_user_model

User = get_user_model()
logger = logging.getLogger(__name__)


class Document(models.Model):
    title = models.CharField(max_length=255)
    file = models.FileField(upload_to="documents/", blank=True, null=True)
    content = models.TextField(blank=True, null=True)

    doc_type = models.CharField(max_length=50, blank=True, null=True)
    category = models.CharField(max_length=50, blank=True, null=True)
    tags = models.JSONField(default=list, blank=True)

    is_active = models.BooleanField(default=True)
    uploaded_by = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, blank=True)

    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def save(self, *args, **kwargs):
        """
        On save, if a file is uploaded but no content is provided,
        attempt to extract text from the file.
        """
        if self.file and not self.content:
            ext = os.path.splitext(self.file.name)[1].lower()
            try:
                self.file.seek(0)  # reset pointer before reading

                if ext == ".txt":
                    self.content = self.file.read().decode("utf-8", errors="ignore")

                elif ext == ".pdf":
                    from PyPDF2 import PdfReader
                    reader = PdfReader(self.file)
                    texts = [page.extract_text() for page in reader.pages if page.extract_text()]
                    self.content = " ".join(texts)

                elif ext == ".docx":
                    import docx
                    doc = docx.Document(self.file)
                    self.content = " ".join([p.text for p in doc.paragraphs if p.text.strip()])

            except Exception as e:
                logger.warning("Failed to extract text from %s: %s", self.file.name, e)

        super().save(*args, **kwargs)

    def __str__(self):
        return self.title or f"Document {self.pk}"
